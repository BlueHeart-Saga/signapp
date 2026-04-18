import os
import json
import uuid
import asyncio
import re
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timezone
from enum import Enum
import base64
from io import BytesIO
import cohere
try:
    from cohere import CohereAPIError
except ImportError:
    try:
        from cohere.errors import ApiError as CohereAPIError
    except ImportError:
        CohereAPIError = Exception

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, Query
from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse
from pydantic import BaseModel, Field, field_validator, ConfigDict
import pypdf
from docx import Document as DocxDocument
from bson import ObjectId

from database import (
    db,  # Your existing database connection
    templates_collection,
    template_fields_collection,
    ai_generation_logs_collection,
    template_versions_collection,
    user_actions_collection
)
from .auth import get_current_user
from .documents import serialize_document, serialize_field_with_recipient
from .fields import normalize_field_value

# Configuration
COHERE_API_KEY = os.getenv("COHERE_API_KEY")

# Initialize Cohere client with robust version detection
cohere_client = None
co_version = None

if COHERE_API_KEY:
    try:
        # Try V2 client first (preferred in newer SDKs)
        if hasattr(cohere, 'ClientV2'):
            cohere_client = cohere.ClientV2(api_key=COHERE_API_KEY)
            co_version = 2
        else:
            # Fallback to V1 client
            cohere_client = cohere.Client(api_key=COHERE_API_KEY)
            co_version = 1
    except Exception as e:
        print(f"[COHERE INIT ERROR] {e}")
        # Try fallback one more time just in case
        try:
            cohere_client = cohere.Client(api_key=COHERE_API_KEY)
            co_version = 1
        except:
            co_version = 0

router = APIRouter(prefix="/api/ai/templates", tags=["AI Template Builder"])

# ========== MODELS ==========
class TemplateType(str, Enum):
    CONTRACT = "contract"
    AGREEMENT = "agreement"
    NDA = "nda"
    PROPOSAL = "proposal"
    INVOICE = "invoice"
    FORM = "form"
    LETTER = "letter"
    RESUME = "resume"
    REPORT = "report"
    MEMO = "memo"
    APPLICATION = "application"
    RECEIPT = "receipt"
    QUOTATION = "quotation"
    ORDER_FORM = "order_form"
    CHECKLIST = "checklist"
    QUESTIONNAIRE = "questionnaire"
    FEEDBACK_FORM = "feedback_form"
    SURVEY = "survey"
    CERTIFICATE = "certificate"
    OTHER = "other"

class FieldType(str, Enum):
    TEXT = "text"
    DATE = "date"
    SIGNATURE = "signature"
    INITIAL = "initial"
    CHECKBOX = "checkbox"
    RADIO = "radio"
    DROPDOWN = "dropdown"
    EMAIL = "email"
    PHONE = "phone"
    NUMBER = "number"
    TEXTAREA = "textarea"
    FILE = "file"
    IMAGE = "image"
    QR_CODE = "qr_code"
    TABLE = "table"
    SECTION = "section"
    HEADER = "header"
    FOOTER = "footer"

class DocumentStyle(str, Enum):
    MODERN = "modern"
    CLASSIC = "classic"
    FORMAL = "formal"
    CASUAL = "casual"
    LEGAL = "legal"
    BUSINESS = "business"
    CREATIVE = "creative"
    MINIMAL = "minimal"

class TemplateFieldSchema(BaseModel):
    id: Optional[str] = None
    name: str
    label: str
    field_type: FieldType
    required: bool = True
    placeholder: Optional[str] = None
    default_value: Optional[str] = None
    options: Optional[List[str]] = None
    x_position: float = Field(ge=0, le=100, default=0)
    y_position: float = Field(ge=0, le=100, default=0)
    width: float = Field(ge=0, le=100, default=0)
    height: float = Field(ge=0, le=100, default=0)
    style: Optional[Dict[str, Any]] = None
    validation_rules: Optional[Dict[str, Any]] = None
    metadata: Optional[Dict[str, Any]] = None

    model_config = ConfigDict(use_enum_values=True)

class AITemplateRequest(BaseModel):
    description: str
    template_type: TemplateType = TemplateType.CONTRACT
    document_style: DocumentStyle = DocumentStyle.MODERN
    fields_to_extract: Optional[List[str]] = None
    language: str = "en"
    tone: str = "professional"
    include_clauses: Optional[List[str]] = None
    custom_instructions: Optional[str] = None
    tags: Optional[List[str]] = None
    generate_html: bool = True
    include_styles: bool = True

    @field_validator('description')
    def description_not_empty(cls, v):
        if not v or len(v.strip()) < 10:
            raise ValueError('Description must be at least 10 characters')
        return v.strip()

class AITemplateResponse(BaseModel):
    template_id: str
    title: str
    content: str
    html_content: str
    fields: List[TemplateFieldSchema]
    summary: str
    suggested_tags: List[str]
    metadata: Dict[str, Any]
    created_at: datetime
    version: int = 1

class DocumentAnalysisResponse(BaseModel):
    document_type: str
    summary: str
    potential_fields: List[TemplateFieldSchema]
    key_clauses: List[str]
    risk_level: str
    recommendations: List[str]
    word_count: int
    page_count: Optional[int] = None
    suggested_template_type: TemplateType

class TemplateToDocumentRequest(BaseModel):
    template_id: str
    filename: Optional[str] = None
    envelope_id: Optional[str] = None
    auto_generate_envelope: bool = True

# ========== HELPER FUNCTIONS ==========
async def _safe_cohere_call(model, messages, response_format):
    retries = 5
    for attempt in range(retries):
        try:
            if co_version == 2:
                # Cohere V2 usage
                response = cohere_client.chat(
                    model=model,
                    messages=messages,
                    response_format={"type": response_format} if response_format == "json_object" else None,
                    temperature=0.3,
                    max_tokens=4000
                )
                return response
            else:
                # Cohere V1 (classic) usage
                # Extract preamble from system message and message from user message
                preamble = messages[0]["content"] if messages[0]["role"] == "system" else ""
                user_msg = messages[-1]["content"] if messages[-1]["role"] == "user" else ""
                
                # Convert history
                history = []
                for msg in messages[1:-1]:
                    history.append({"role": "USER" if msg["role"] == "user" else "CHATBOT", "message": msg["content"]})
                
                response = cohere_client.chat(
                    model=model,
                    message=user_msg,
                    preamble=preamble,
                    chat_history=history,
                    # Classic V1 doesn't consistently support response_format keyword in all SDK versions.
                    temperature=0.3,
                    max_tokens=4000
                )
                return response

        except CohereAPIError as e:
            if e.status_code == 429:
                wait = (attempt + 1) * 4.0
                print(f"[Cohere 429] Rate limited (Attempt {attempt + 1}/{retries}). Retrying in {wait}s...")
                await asyncio.sleep(wait)
            else:
                print(f"[Cohere ERROR] {e}")
                raise HTTPException(status_code=500, detail=f"Cohere service error: {str(e)}")
        except Exception as e:
            if isinstance(e, HTTPException):
                raise e
            print(f"[AI ERROR] {e}")
            raise HTTPException(status_code=500, detail=f"AI call failed: {str(e)}")
    
    raise HTTPException(
        status_code=429,
        detail="AI service is currently experiencing high traffic. Please try again in a few minutes."
    )
    
async def analyze_with_ai(document_text: str, instructions: str, response_format="json_object") -> Dict:
    if not cohere_client:
        raise HTTPException(status_code=500, detail="AI client not configured")

    try:
        json_req = " STRICTLY return ONLY valid JSON." if response_format == "json_object" else ""
        messages = [
            {"role": "system", "content": f"You are a professional document template generator and legal expert.{json_req}"},
            {"role": "user", "content": f"{instructions}\n\n{document_text}"}
        ]

        model = "command-r-08-2024" if len(document_text) < 4000 else "command-r-plus-08-2024"

        # SAFE CALL WRAPPER
        response = await _safe_cohere_call(model, messages, response_format)

        if co_version == 2:
            content = response.message.content[0].text
        else:
            content = response.text

        if response_format == "json_object":
            # Extract the JSON object from the response string
            try:
                data = json.loads(content)
            except json.JSONDecodeError:
                # Fallback: Find the first { and last } to handle preamble/postamble
                start = content.find('{')
                end = content.rfind('}') + 1
                if start != -1 and end != 0:
                    data = json.loads(content[start:end])
                else:
                    raise
            
            # Post-process content to handle Cohere's tendency to wrap HTML tags in braces
            if isinstance(data, dict) and "content" in data:
                c = data["content"]
                # Fix triple braces: {{{h1}}} -> <h1>, {{{\/h1}}} -> </h1>
                c = re.sub(r'{{{([a-z1-6]+)}}}', r'<\1>', c)
                c = re.sub(r'{{{\/([a-z1-6]+)}}}', r'</\1>', c)
                # Fix double braces: {{h1}} -> <h1>, {{\/h1}} -> </h1>
                c = re.sub(r'{{([a-z1-6]+)}}', r'<\1>', c)
                c = re.sub(r'{{\/([a-z1-6]+)}}', r'</\1>', c)
                # Ensure fields stay as {{field_name}} - fix any over-conversion
                # Re-check field format: {{field}} is already the standard
                data["content"] = c
                
            return data

        return content

    except HTTPException:
        # Let FastAPI HTTPExceptions through (like our 429)
        raise
    except Exception as e:
        print(f"[ANALYSIS ERROR] {e}")
        raise HTTPException(status_code=500, detail=f"OpenAI analysis failed: {str(e)}")

async def extract_document_text(file: UploadFile) -> Tuple[str, int]:
    """Extract text from uploaded document"""
    content = await file.read()
    filename = file.filename or "document"
    file_extension = filename.split('.')[-1].lower() if '.' in filename else ''
    
    try:
        page_count = 1
        
        if file_extension == 'pdf':
            pdf_reader = pypdf.PdfReader(BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n\n"
            page_count = len(pdf_reader.pages)
            return text.strip(), page_count
        
        elif file_extension in ['doc', 'docx']:
            try:
                doc = DocxDocument(BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\t"
                        text += "\n"
                page_count = max(1, len(text.split()) // 500)
                return text.strip(), page_count
            except:
                raise HTTPException(status_code=400, detail="DOC format not fully supported. Please convert to DOCX or PDF.")
        
        elif file_extension in ['txt', 'md', 'rtf', 'html', 'htm']:
            try:
                text = content.decode('utf-8')
                page_count = max(1, len(text.split()) // 500)
                return text.strip(), page_count
            except UnicodeDecodeError:
                text = content.decode('latin-1')
                page_count = max(1, len(text.split()) // 500)
                return text.strip(), page_count
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {file_extension}")
                
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")

def get_default_field_dimensions(field_type: str) -> Tuple[float, float]:
    """Get default width and height for field types"""
    dimensions = {
        "text": (30.0, 5.0),
        "email": (35.0, 5.0),
        "phone": (30.0, 5.0),
        "number": (25.0, 5.0),
        "date": (25.0, 5.0),
        "textarea": (80.0, 15.0),
        "signature": (60.0, 20.0),
        "initial": (20.0, 8.0),
        "dropdown": (35.0, 6.0),
        "checkbox": (15.0, 4.0),
        "radio": (15.0, 4.0),
    }
    return dimensions.get(field_type, (30.0, 5.0))

def normalize_field_names(content: str) -> Tuple[str, List[str]]:
    """Normalize field names in content and extract them"""
    # Define field mappings for inconsistent field names
    field_mappings = {
        "(client_)": "(client_name)",
        "(client": "(client_name)",
        "(client_name": "(client_name)",
        "(service_provider": "(service_provider_name)",
        "(special_amount)": "(deposit_amount)",
        "(total_date)": "(completion_date)",
        "(signature_date)": "(client_signature_date)",
        "[signature_date]": "(service_provider_signature_date)",
        "(governing_law_state)": "(governing_state)",
        "(jurisdiction)": "(court_jurisdiction)"
    }
    
    # Apply mappings
    normalized_content = content
    for old, new in field_mappings.items():
        normalized_content = normalized_content.replace(old, new)
    
    # Extract field names using multiple patterns
    field_names = set()
    patterns = [
        r'\{\{([^{}]+)\}\}',      # {{field_name}}
        r'\{([^{}]+)\}',          # {field_name}
        r'\(([^()]+)\)',          # (field_name)
        r'\[([^\[\]]+)\]',        # [field_name]
        r'<([^<>]+)>',            # <field_name>
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, normalized_content)
        for match in matches:
            field_name = match.strip()
            if field_name:
                # Convert to snake_case if needed
                field_name = re.sub(r'[-\s]+', '_', field_name).lower()
                field_names.add(field_name)
    
    # Replace all patterns with standardized {{field_name}} format
    for field_name in field_names:
        # Replace all possible formats with standardized {{field_name}}
        patterns_to_replace = [
            f"{{{{{field_name}}}}}",
            f"{{{field_name}}}",
            f"({field_name})",
            f"[{field_name}]",
            f"<{field_name}>",
        ]
        
        for pattern in patterns_to_replace:
            if pattern in normalized_content:
                normalized_content = normalized_content.replace(pattern, f"{{{{{field_name}}}}}")
    
    return normalized_content, list(field_names)

def create_edit_mode_html(content: str, fields: List[Dict]) -> str:
    """Create edit mode HTML with interactive placeholders"""
    
    # Normalize content first
    normalized_content, detected_field_names = normalize_field_names(content)
    
    # Create a mapping of field names to field data for quick lookup
    field_map = {field.get("name", ""): field for field in fields}
    
    # Replace placeholders with styled spans
    edit_content = normalized_content
    
    for field_name in detected_field_names:
        field = field_map.get(field_name, {})
        field_type = field.get("field_type") or field.get("type", "text")
        field_label = field.get("label", field_name.replace("_", " ").title())
        required = field.get("required", True)
        
        # Create styled placeholder with all field info
        styled_placeholder = f'''
        <span class="field-placeholder" data-field-name="{field_name}" data-field-type="{field_type}" data-required="{str(required).lower()}">
            <span class="placeholder-wrapper">
                <span class="placeholder-brackets">{{{{</span>
                <span class="placeholder-name">{field_label}</span>
                <span class="placeholder-brackets">}}}}</span>
                <span class="field-type-indicator">{field_type}</span>
                {'' if not required else '<span class="required-indicator">*</span>'}
            </span>
        </span>
        '''
        
        # Replace the placeholder
        placeholder_pattern = f"{{{{{field_name}}}}}"
        edit_content = edit_content.replace(placeholder_pattern, styled_placeholder)
    
    # HTML template for edit mode
    html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Template Edit Mode</title>
        <style>
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
                line-height: 1.6;
                color: #2c3e50;
                background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                padding: 20px;
                min-height: 100vh;
            }}
            
            .edit-container {{
                max-width: 850px;
                margin: 0 auto;
                background: white;
                border-radius: 12px;
                box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
                padding: 50px;
                position: relative;
                min-height: 100vh;
                background-image: 
                    linear-gradient(90deg, rgba(0, 0, 0, 0.03) 1px, transparent 1px),
                    linear-gradient(180deg, rgba(0, 0, 0, 0.03) 1px, transparent 1px);
                background-size: 25px 25px;
                white-space: pre-wrap;
                overflow-wrap: break-word;
                font-size: 16px;
                line-height: 1.8;
            }}
            
            .edit-container h1 {{
                font-size: 2.5em;
                margin-bottom: 30px;
                color: #1a237e;
                border-bottom: 3px solid #3f51b5;
                padding-bottom: 15px;
            }}
            
            .edit-container h2 {{
                font-size: 2em;
                margin: 25px 0 15px 0;
                color: #283593;
            }}
            
            .edit-container h3 {{
                font-size: 1.5em;
                margin: 20px 0 12px 0;
                color: #3949ab;
            }}
            
            .edit-container p {{
                margin: 12px 0;
                text-align: justify;
            }}
            
            .field-placeholder {{
                display: inline-block;
                background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
                border: 2px solid #2196f3;
                border-radius: 8px;
                padding: 8px 12px;
                margin: 4px 2px;
                cursor: pointer;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
                font-family: 'SF Mono', Monaco, 'Cascadia Code', monospace;
                vertical-align: middle;
                position: relative;
                box-shadow: 0 2px 5px rgba(33, 150, 243, 0.2);
            }}
            
            .field-placeholder:hover {{
                background: linear-gradient(135deg, #bbdefb 0%, #90caf9 100%);
                border-color: #1976d2;
                transform: translateY(-2px) scale(1.02);
                box-shadow: 0 5px 15px rgba(33, 150, 243, 0.3);
                z-index: 10;
            }}
            
            .field-placeholder:active {{
                transform: translateY(-1px) scale(1.01);
            }}
            
            .placeholder-wrapper {{
                display: flex;
                align-items: center;
                gap: 6px;
            }}
            
            .placeholder-brackets {{
                color: #1565c0;
                font-weight: 700;
                font-size: 0.9em;
            }}
            
            .placeholder-name {{
                color: #0d47a1;
                font-weight: 600;
                font-size: 0.95em;
                padding: 0 2px;
            }}
            
            .field-type-indicator {{
                background: #1a237e;
                color: white;
                font-size: 0.7em;
                padding: 2px 6px;
                border-radius: 4px;
                font-weight: 500;
                text-transform: uppercase;
                letter-spacing: 0.5px;
                margin-left: 4px;
            }}
            
            .required-indicator {{
                color: #d32f2f;
                font-weight: bold;
                font-size: 1.2em;
                margin-left: 2px;
            }}
            
            .field-info-panel {{
                position: fixed;
                top: 20px;
                right: 20px;
                background: white;
                border: 1px solid #e0e0e0;
                border-radius: 10px;
                padding: 20px;
                width: 300px;
                max-height: 400px;
                box-shadow: 0 8px 25px rgba(0, 0, 0, 0.15);
                display: none;
                z-index: 1000;
                overflow-y: auto;
                backdrop-filter: blur(10px);
                background: rgba(255, 255, 255, 0.95);
            }}
            
            .field-info-panel h4 {{
                margin-bottom: 15px;
                color: #1a237e;
                font-size: 1.2em;
                border-bottom: 2px solid #3f51b5;
                padding-bottom: 8px;
                display: flex;
                align-items: center;
                gap: 8px;
            }}
            
            .field-info-item {{
                margin: 10px 0;
                font-size: 14px;
                display: flex;
                align-items: flex-start;
            }}
            
            .field-info-label {{
                font-weight: 600;
                color: #546e7a;
                min-width: 100px;
                margin-right: 10px;
            }}
            
            .field-info-value {{
                color: #263238;
                flex: 1;
                word-break: break-word;
            }}
            
            .edit-toolbar {{
                position: fixed;
                bottom: 30px;
                right: 30px;
                background: white;
                border-radius: 50px;
                padding: 15px 25px;
                box-shadow: 0 5px 20px rgba(0, 0, 0, 0.15);
                display: flex;
                gap: 15px;
                z-index: 100;
            }}
            
            .edit-button {{
                background: #3f51b5;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 25px;
                cursor: pointer;
                font-weight: 500;
                transition: all 0.3s ease;
            }}
            
            .edit-button:hover {{
                background: #303f9f;
                transform: translateY(-2px);
            }}
            
            @media (max-width: 768px) {{
                .edit-container {{
                    padding: 25px;
                    font-size: 15px;
                }}
                
                .field-info-panel {{
                    position: fixed;
                    top: auto;
                    bottom: 0;
                    left: 0;
                    right: 0;
                    width: 100%;
                    max-height: 50vh;
                    border-radius: 15px 15px 0 0;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="edit-container">
            {edit_content}
        </div>
        
        <div id="fieldInfo" class="field-info-panel"></div>
        
        <div class="edit-toolbar">
            <button class="edit-button" onclick="saveTemplate()">💾 Save</button>
            <button class="edit-button" onclick="previewTemplate()">👁️ Preview</button>
            <button class="edit-button" onclick="convertToDocument()">📄 Convert to Document</button>
        </div>
        
        <script>
            let activeFieldName = '';
            const templateId = window.location.pathname.split('/').pop().replace('/edit-mode', '');
            
            document.addEventListener('DOMContentLoaded', function() {{
                const fieldInfo = document.getElementById('fieldInfo');
                
                // Add hover and click effects to placeholders
                document.querySelectorAll('.field-placeholder').forEach(placeholder => {{
                    placeholder.addEventListener('mouseenter', function(e) {{
                        const fieldName = this.dataset.fieldName;
                        const fieldType = this.dataset.fieldType;
                        const required = this.dataset.required === 'true';
                        activeFieldName = fieldName;
                        
                        // Show field info panel
                        fieldInfo.innerHTML = `
                            <h4>
                                <span style="background: #e3f2fd; padding: 4px 8px; border-radius: 4px;">{{</span>
                                ${{fieldName}}
                                <span style="background: #e3f2fd; padding: 4px 8px; border-radius: 4px;">}}</span>
                            </h4>
                            <div class="field-info-item">
                                <span class="field-info-label">Field Label:</span>
                                <span class="field-info-value">${{this.querySelector('.placeholder-name').textContent}}</span>
                            </div>
                            <div class="field-info-item">
                                <span class="field-info-label">Field Type:</span>
                                <span class="field-info-value">
                                    <span style="background: #1a237e; color: white; padding: 2px 8px; border-radius: 4px; font-size: 12px;">
                                        ${{fieldType.toUpperCase()}}
                                    </span>
                                </span>
                            </div>
                            <div class="field-info-item">
                                <span class="field-info-label">Required:</span>
                                <span class="field-info-value">
                                    <span style="color: ${{required ? '#d32f2f' : '#4caf50'}}; font-weight: bold;">
                                        ${{required ? 'Yes' : 'No'}}
                                    </span>
                                </span>
                            </div>
                            <div class="field-info-item">
                                <span class="field-info-label">Placeholder:</span>
                                <span class="field-info-value">"Enter ${{fieldName.replace(/_/g, ' ')}}"</span>
                            </div>
                            <div class="field-info-item">
                                <span class="field-info-label">Actions:</span>
                                <span class="field-info-value">
                                    <button onclick="editField('${{fieldName}}')" style="background: #2196f3; color: white; border: none; padding: 5px 10px; border-radius: 4px; margin-right: 5px; cursor: pointer;">Edit</button>
                                    <button onclick="deleteField('${{fieldName}}')" style="background: #f44336; color: white; border: none; padding: 5px 10px; border-radius: 4px; cursor: pointer;">Delete</button>
                                </span>
                            </div>
                        `;
                        fieldInfo.style.display = 'block';
                        
                        // Position panel near cursor but not off-screen
                        const x = e.clientX + 20;
                        const y = e.clientY - 50;
                        const panelWidth = fieldInfo.offsetWidth;
                        const panelHeight = fieldInfo.offsetHeight;
                        const windowWidth = window.innerWidth;
                        const windowHeight = window.innerHeight;
                        
                        fieldInfo.style.left = (x + panelWidth > windowWidth ? windowWidth - panelWidth - 20 : x) + 'px';
                        fieldInfo.style.top = (y + panelHeight > windowHeight ? windowHeight - panelHeight - 20 : y) + 'px';
                    }});
                    
                    placeholder.addEventListener('mouseleave', function(e) {{
                        // Only hide if mouse is not over the info panel
                        if (!fieldInfo.matches(':hover')) {{
                            fieldInfo.style.display = 'none';
                        }}
                    }});
                    
                    placeholder.addEventListener('click', function(e) {{
                        const fieldName = this.dataset.fieldName;
                        editField(fieldName);
                    }});
                }});
                
                // Keep info panel visible when hovering over it
                fieldInfo.addEventListener('mouseenter', function() {{
                    this.style.display = 'block';
                }});
                
                fieldInfo.addEventListener('mouseleave', function() {{
                    this.style.display = 'none';
                }});
            }});
            
            function editField(fieldName) {{
                alert(`Opening editor for field: ${{fieldName}}\\n\\nIn a real implementation, this would open a modal with field configuration options.`);
            }}
            
            function deleteField(fieldName) {{
                if (confirm(`Are you sure you want to delete field: ${{fieldName}}?`)) {{
                    alert(`Field ${{fieldName}} deleted (simulated)`);
                    // In real implementation, make API call to delete field
                }}
            }}
            
            function saveTemplate() {{
                alert('Template saved! (simulated)');
                // In real implementation, make API call to save template
            }}
            
            function previewTemplate() {{
                // Switch to preview mode
                window.location.href = window.location.href.replace('/edit-mode', '/preview-mode');
            }}
            
            function convertToDocument() {{
                // Convert template to document
                fetch(`/api/ai/templates/${{templateId}}/convert-to-document`, {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json',
                        'Authorization': localStorage.getItem('token') ? 'Bearer ' + localStorage.getItem('token') : ''
                    }},
                    body: JSON.stringify({{
                        template_id: templateId,
                        filename: 'Converted Document'
                    }})
                }})
                .then(response => response.json())
                .then(data => {{
                    if (data.document_id) {{
                        alert(`Document created successfully!\\nDocument ID: ${{data.document_id}}\\n\\nRedirecting to document...`);
                        window.location.href = `/documents/${{data.document_id}}`;
                    }} else {{
                        alert('Failed to create document: ' + (data.detail || 'Unknown error'));
                    }}
                }})
                .catch(error => {{
                    alert('Error: ' + error.message);
                }});
            }}
        </script>
    </body>
    </html>
    '''
    
    return html

def create_preview_mode_html(content: str, fields: List[Dict], style: str = "modern") -> str:
    """Create preview mode HTML with actual form fields"""
    
    # Normalize content first
    normalized_content, detected_field_names = normalize_field_names(content)
    
    # Create field mapping for quick lookup
    field_map = {field.get("name", ""): field for field in fields}
    
    # Replace placeholders with actual form fields
    preview_content = normalized_content
    
    for field_name in detected_field_names:
        field = field_map.get(field_name, {})
        field_type = field.get("field_type") or field.get("type", "text")
        field_label = field.get("label", field_name.replace("_", " ").title())
        placeholder = field.get("placeholder", f"Enter {field_label}")
        required = field.get("required", False)
        default_value = field.get("default_value", "")
        
        # Determine appropriate field width based on content
        field_width = "auto"
        if field_type in ["text", "email", "phone", "number"]:
            field_width = "200px"
        elif field_type == "date":
            field_width = "180px"
        elif field_type == "textarea":
            field_width = "400px"
        elif field_type in ["signature", "initial"]:
            field_width = "300px"
        
        # Create appropriate HTML for the field type
        field_html = ""
        field_id = f"field_{field_name}"
        
        if field_type == "text":
            field_html = f'''
            <input type="text" 
                   id="{field_id}" 
                   name="{field_name}" 
                   placeholder="{placeholder}" 
                   value="{default_value}"
                   class="form-field {field_type}-field"
                   {"" if not required else "required"}
                   style="width: {field_width};">
            '''
        elif field_type == "email":
            field_html = f'''
            <input type="email" 
                   id="{field_id}" 
                   name="{field_name}" 
                   placeholder="{placeholder}" 
                   value="{default_value}"
                   class="form-field {field_type}-field"
                   {"" if not required else "required"}
                   style="width: {field_width};">
            '''
        elif field_type == "date":
            field_html = f'''
            <input type="date" 
                   id="{field_id}" 
                   name="{field_name}" 
                   value="{default_value}"
                   class="form-field {field_type}-field"
                   {"" if not required else "required"}
                   style="width: {field_width};">
            '''
        elif field_type == "number":
            field_html = f'''
            <input type="number" 
                   id="{field_id}" 
                   name="{field_name}" 
                   placeholder="{placeholder}" 
                   value="{default_value}"
                   class="form-field {field_type}-field"
                   {"" if not required else "required"}
                   style="width: {field_width};">
            '''
        elif field_type == "phone":
            field_html = f'''
            <input type="tel" 
                   id="{field_id}" 
                   name="{field_name}" 
                   placeholder="{placeholder}" 
                   value="{default_value}"
                   class="form-field {field_type}-field"
                   {"" if not required else "required"}
                   pattern="[0-9\\s\\-\\+\\()]*"
                   style="width: {field_width};">
            '''
        elif field_type == "textarea":
            field_html = f'''
            <textarea 
                id="{field_id}" 
                name="{field_name}" 
                placeholder="{placeholder}" 
                class="form-field {field_type}-field"
                {"" if not required else "required"}
                rows="4"
                style="width: {field_width}; resize: vertical;">{default_value}</textarea>
            '''
        elif field_type == "signature":
            field_html = f'''
            <div class="signature-container" data-field-name="{field_name}">
                <div class="signature-field" id="{field_id}">
                    <canvas class="signature-canvas"></canvas>
                    <div class="signature-placeholder">{placeholder or "Click to sign here"}</div>
                    <input type="hidden" name="{field_name}" value="{default_value}">
                </div>
                <div class="signature-controls">
                    <button type="button" class="signature-clear">Clear</button>
                </div>
            </div>
            '''
        elif field_type == "dropdown":
            options = field.get("options", ["Option 1", "Option 2", "Option 3"])
            options_html = ""
            for option in options:
                selected = "selected" if option == default_value else ""
                options_html += f'<option value="{option}" {selected}>{option}</option>'
            field_html = f'''
            <select id="{field_id}" 
                    name="{field_name}" 
                    class="form-field {field_type}-field"
                    {"" if not required else "required"}
                    style="width: {field_width};">
                <option value="">{placeholder or "Select an option"}</option>
                {options_html}
            </select>
            '''
        else:
            # Default to text field
            field_html = f'''
            <input type="text" 
                   id="{field_id}" 
                   name="{field_name}" 
                   placeholder="{placeholder}" 
                   value="{default_value}"
                   class="form-field text-field"
                   {"" if not required else "required"}
                   style="width: {field_width};">
            '''
        
        # Wrap in a container with label
        field_container = f'''
        <div class="field-container" data-field-name="{field_name}" data-field-type="{field_type}">
            <div class="field-label-container">
                <label for="{field_id}" class="field-label">
                    {field_label}
                    {'' if not required else '<span class="required-star">*</span>'}
                </label>
            </div>
            <div class="field-input-container">
                {field_html}
            </div>
            <div class="field-hint">Click to fill this field</div>
        </div>
        '''
        
        # Replace the placeholder with field container
        placeholder_pattern = f"{{{{{field_name}}}}}"
        preview_content = preview_content.replace(placeholder_pattern, field_container)
    
    # Style definitions based on selected style
    style_css = {
        "modern": """
            .preview-container {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                line-height: 1.8;
                color: #333;
                max-width: 900px;
                margin: 0 auto;
                padding: 50px;
                background: white;
                border-radius: 16px;
                box-shadow: 0 15px 35px rgba(0, 0, 0, 0.1);
                position: relative;
            }
            
            .preview-container h1 {
                font-size: 2.8em;
                margin-bottom: 40px;
                color: #1a237e;
                text-align: center;
                border-bottom: 3px solid #3f51b5;
                padding-bottom: 20px;
            }
            
            .preview-container h2 {
                font-size: 2.2em;
                margin: 35px 0 20px 0;
                color: #283593;
                border-left: 5px solid #3f51b5;
                padding-left: 15px;
            }
            
            .preview-container h3 {
                font-size: 1.8em;
                margin: 25px 0 15px 0;
                color: #3949ab;
            }
            
            .preview-container p {
                margin: 15px 0;
                text-align: justify;
                line-height: 1.9;
            }
            
            .field-container {
                margin: 20px 0;
                padding: 20px;
                background: linear-gradient(135deg, #f5f7fa 0%, #e4e8f0 100%);
                border-radius: 12px;
                border: 2px solid #e0e0e0;
                transition: all 0.3s ease;
                position: relative;
            }
            
            .field-container:hover {
                border-color: #2196f3;
                box-shadow: 0 5px 20px rgba(33, 150, 243, 0.15);
                transform: translateY(-2px);
            }
            
            .field-label-container {
                margin-bottom: 10px;
            }
            
            .field-label {
                font-weight: 600;
                color: #1a237e;
                font-size: 1.1em;
                display: block;
                margin-bottom: 8px;
            }
            
            .required-star {
                color: #d32f2f;
                margin-left: 4px;
                font-weight: bold;
            }
            
            .field-input-container {
                margin: 10px 0;
            }
            
            .form-field {
                padding: 12px 16px;
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                font-size: 16px;
                transition: all 0.3s ease;
                background: white;
                width: 100%;
                font-family: inherit;
            }
            
            .form-field:focus {
                outline: none;
                border-color: #2196f3;
                box-shadow: 0 0 0 3px rgba(33, 150, 243, 0.2);
            }
            
            textarea.form-field {
                min-height: 120px;
                line-height: 1.5;
                resize: vertical;
            }
            
            .signature-container {
                border: 2px dashed #90a4ae;
                border-radius: 8px;
                padding: 20px;
                background: white;
                cursor: pointer;
                position: relative;
                min-height: 150px;
                transition: all 0.3s ease;
            }
            
            .signature-container:hover {
                border-color: #2196f3;
                background: #f8fdff;
            }
            
            .signature-field {
                position: relative;
                min-height: 120px;
            }
            
            .signature-canvas {
                border: 1px solid #ddd;
                border-radius: 4px;
                background: white;
                width: 100%;
                height: 120px;
                display: none;
            }
            
            .signature-placeholder {
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                color: #78909c;
                font-style: italic;
                font-size: 1.1em;
                text-align: center;
                width: 100%;
            }
            
            .signature-controls {
                margin-top: 10px;
                text-align: right;
            }
            
            .signature-clear {
                background: #f44336;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                cursor: pointer;
                font-size: 14px;
            }
            
            .signature-clear:hover {
                background: #d32f2f;
            }
            
            .field-hint {
                font-size: 12px;
                color: #78909c;
                margin-top: 8px;
                font-style: italic;
                display: none;
            }
            
            .field-container:hover .field-hint {
                display: block;
            }
        """,
        "classic": """
            .preview-container {
                
                line-height: 1.7;
                color: #000;
                max-width: 800px;
                margin: 0 auto;
                padding: 40px;
                background: white;
                border: 2px solid #8b7355;
                position: relative;
            }
            
            .preview-container h1 {
                font-size: 2.5em;
                margin-bottom: 30px;
                color: #000;
                text-align: center;
                border-bottom: 2px solid #8b7355;
                padding-bottom: 15px;
                font-weight: bold;
            }
            
            .preview-container h2 {
                font-size: 2em;
                margin: 25px 0 15px 0;
                color: #333;
                font-weight: bold;
            }
            
            .preview-container h3 {
                font-size: 1.6em;
                margin: 20px 0 12px 0;
                color: #555;
                font-weight: bold;
            }
            
            .field-container {
                margin: 15px 0;
                padding: 15px;
                background: #f9f7f3;
                border: 1px solid #d4c4a8;
                border-left: 4px solid #8b7355;
            }
            
            .form-field {
                padding: 10px 12px;
                border: 1px solid #8b7355;
                font-size: 15px;
                background: white;
                width: 100%;
                font-family: inherit;
            }
            
            .signature-container {
                border: 2px dashed #8b7355;
                padding: 15px;
                background: white;
                min-height: 130px;
            }
        """,
        "legal": """
            .preview-container {
                font-family: 'Georgia', serif;
                line-height: 1.9;
                color: #222;
                max-width: 850px;
                margin: 0 auto;
                padding: 45px;
                background: white;
                border: 2px solid #8b0000;
                position: relative;
                counter-reset: section;
            }
            
            .preview-container h1 {
                font-size: 2.6em;
                margin-bottom: 35px;
                color: #8b0000;
                text-align: center;
                text-transform: uppercase;
                letter-spacing: 1px;
                border-bottom: 3px solid #8b0000;
                padding-bottom: 20px;
            }
            
            .preview-container h2 {
                font-size: 2em;
                margin: 30px 0 18px 0;
                color: #8b0000;
                border-left: 4px solid #8b0000;
                padding-left: 15px;
                counter-increment: section;
            }
            
            .preview-container h2:before {
                content: counter(section) ". ";
                font-weight: bold;
            }
            
            .field-container {
                margin: 18px 0;
                padding: 18px;
                background: #fff5f5;
                border: 2px solid #8b0000;
                border-radius: 4px;
            }
            
            .form-field {
                padding: 11px 14px;
                border: 2px solid #8b0000;
                font-size: 15px;
                background: #fffafa;
                width: 100%;
                font-family: inherit;
            }
            
            .signature-container {
                border: 3px double #8b0000;
                padding: 20px;
                background: #fffafa;
                min-height: 140px;
            }
        """
    }
    
    selected_style = style_css.get(style, style_css["modern"])
    
    # HTML template for preview mode
    html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Template Preview</title>
        <style>
            * {{
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }}
            
            body {{
                background: linear-gradient(135deg, #e0eafc 0%, #cfdef3 100%);
                padding: 30px;
                min-height: 100vh;
            }}
            
            {selected_style}
            
            .preview-header {{
                text-align: center;
                margin-bottom: 40px;
                color: #546e7a;
                font-size: 14px;
            }}
            
            .preview-footer {{
                margin-top: 50px;
                padding-top: 20px;
                border-top: 2px solid #e0e0e0;
                text-align: center;
                color: #78909c;
                font-size: 13px;
            }}
            
            .form-actions {{
                margin-top: 40px;
                text-align: center;
                padding: 20px;
                background: #f5f7fa;
                border-radius: 10px;
                border: 2px dashed #b0bec5;
            }}
            
            .submit-button {{
                background: linear-gradient(135deg, #4CAF50 0%, #45a049 100%);
                color: white;
                border: none;
                padding: 15px 40px;
                border-radius: 8px;
                font-size: 16px;
                font-weight: 600;
                cursor: pointer;
                transition: all 0.3s ease;
                box-shadow: 0 4px 15px rgba(76, 175, 80, 0.3);
            }}
            
            .submit-button:hover {{
                transform: translateY(-2px);
                box-shadow: 0 6px 20px rgba(76, 175, 80, 0.4);
                background: linear-gradient(135deg, #45a049 0%, #3d8b40 100%);
            }}
            
            .edit-mode-button {{
                background: #2196f3;
                color: white;
                border: none;
                padding: 12px 25px;
                border-radius: 6px;
                cursor: pointer;
                margin-left: 15px;
                transition: all 0.3s ease;
            }}
            
            .edit-mode-button:hover {{
                background: #1976d2;
                transform: translateY(-1px);
            }}
            
            .convert-button {{
                background: #673ab7;
                color: white;
                border: none;
                padding: 12px 25px;
                border-radius: 6px;
                cursor: pointer;
                margin-left: 15px;
                transition: all 0.3s ease;
            }}
            
            .convert-button:hover {{
                background: #5e35b1;
                transform: translateY(-1px);
            }}
            
            @media (max-width: 768px) {{
                body {{
                    padding: 15px;
                }}
                
                .preview-container {{
                    padding: 25px;
                }}
                
                .form-field {{
                    font-size: 15px;
                }}
                
                .submit-button, .edit-mode-button, .convert-button {{
                    width: 100%;
                    margin: 10px 0;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="preview-header">
            <h2>📄 Document Preview Mode</h2>
            <p>Fill in all required fields (marked with *) and submit the form</p>
        </div>
        
        <div class="preview-container">
            {preview_content}
            
            <div class="form-actions">
                <button type="button" class="submit-button" onclick="submitForm()">Submit Form</button>
                <button type="button" class="edit-mode-button" onclick="switchToEditMode()">✏️ Edit Template</button>
                <button type="button" class="convert-button" onclick="convertToDocument()">📄 Convert to Document</button>
            </div>
        </div>
        
        <div class="preview-footer">
            <p>Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")} | Total Fields: {len(detected_field_names)}</p>
        </div>
        
        <script>
            const templateId = window.location.pathname.split('/').pop().replace('/preview-mode', '');
            
            document.addEventListener('DOMContentLoaded', function() {{
                // Initialize signature fields
                document.querySelectorAll('.signature-container').forEach(container => {{
                    const canvas = container.querySelector('.signature-canvas');
                    const placeholder = container.querySelector('.signature-placeholder');
                    const hiddenInput = container.querySelector('input[type="hidden"]');
                    const clearBtn = container.querySelector('.signature-clear');
                    
                    container.addEventListener('click', function(e) {{
                        if (e.target !== clearBtn) {{
                            canvas.style.display = 'block';
                            placeholder.style.display = 'none';
                            
                            canvas.width = canvas.offsetWidth;
                            canvas.height = canvas.offsetHeight;
                            
                            const ctx = canvas.getContext('2d');
                            let isDrawing = false;
                            let lastX = 0;
                            let lastY = 0;
                            
                            function draw(e) {{
                                if (!isDrawing) return;
                                
                                ctx.beginPath();
                                ctx.moveTo(lastX, lastY);
                                ctx.lineTo(e.offsetX, e.offsetY);
                                ctx.strokeStyle = '#000';
                                ctx.lineWidth = 2;
                                ctx.lineCap = 'round';
                                ctx.lineJoin = 'round';
                                ctx.stroke();
                                
                                [lastX, lastY] = [e.offsetX, e.offsetY];
                                
                                // Update hidden input with data URL
                                hiddenInput.value = canvas.toDataURL();
                            }}
                            
                            canvas.addEventListener('mousedown', (e) => {{
                                isDrawing = true;
                                [lastX, lastY] = [e.offsetX, e.offsetY];
                            }});
                            
                            canvas.addEventListener('mousemove', draw);
                            canvas.addEventListener('mouseup', () => isDrawing = false);
                            canvas.addEventListener('mouseout', () => isDrawing = false);
                        }}
                    }});
                    
                    if (clearBtn) {{
                        clearBtn.addEventListener('click', function(e) {{
                            e.stopPropagation();
                            const ctx = canvas.getContext('2d');
                            ctx.clearRect(0, 0, canvas.width, canvas.height);
                            canvas.style.display = 'none';
                            placeholder.style.display = 'block';
                            hiddenInput.value = '';
                        }});
                    }}
                }});
                
                // Add focus styles to all form fields
                document.querySelectorAll('.form-field').forEach(field => {{
                    field.addEventListener('focus', function() {{
                        this.parentElement.parentElement.style.borderColor = '#2196f3';
                        this.parentElement.parentElement.style.boxShadow = '0 5px 20px rgba(33, 150, 243, 0.2)';
                    }});
                    
                    field.addEventListener('blur', function() {{
                        this.parentElement.parentElement.style.borderColor = '#e0e0e0';
                        this.parentElement.parentElement.style.boxShadow = 'none';
                    }});
                }});
                
                // Validate required fields on input
                document.querySelectorAll('.form-field[required]').forEach(field => {{
                    field.addEventListener('input', function() {{
                        if (this.value.trim()) {{
                            this.style.borderColor = '#4CAF50';
                        }} else {{
                            this.style.borderColor = '#f44336';
                        }}
                    }});
                }});
            }});
            
            function submitForm() {{
                // Collect all form data
                const formData = {{}};
                let isValid = true;
                
                document.querySelectorAll('.form-field').forEach(field => {{
                    if (field.tagName === 'INPUT' || field.tagName === 'TEXTAREA' || field.tagName === 'SELECT') {{
                        const fieldName = field.name || field.id.replace('field_', '');
                        formData[fieldName] = field.value;
                        
                        // Validate required fields
                        if (field.hasAttribute('required') && !field.value.trim()) {{
                            field.style.borderColor = '#f44336';
                            isValid = false;
                        }}
                    }}
                }});
                
                if (!isValid) {{
                    alert('⚠️ Please fill in all required fields (marked with *)');
                    return;
                }}
                
                // Show submission confirmation
                const fieldCount = Object.keys(formData).length;
                alert(`Form submitted successfully!\\n\\nFields filled: ${{fieldCount}}\\n\\nData collected:\\n${{JSON.stringify(formData, null, 2)}}`);
                
                // In real implementation, send data to server
                console.log('Form data:', formData);
            }}
            
            function switchToEditMode() {{
                // Switch to edit mode
                window.location.href = window.location.href.replace('/preview-mode', '/edit-mode');
            }}
            
            function convertToDocument() {{
                // Convert template to document
                fetch(`/api/ai/templates/${{templateId}}/convert-to-document`, {{
                    method: 'POST',
                    headers: {{
                        'Content-Type': 'application/json',
                        'Authorization': localStorage.getItem('token') ? 'Bearer ' + localStorage.getItem('token') : ''
                    }},
                    body: JSON.stringify({{
                        template_id: templateId,
                        filename: 'Converted Document'
                    }})
                }})
                .then(response => response.json())
                .then(data => {{
                    if (data.document_id) {{
                        alert(`Document created successfully!\\nDocument ID: ${{data.document_id}}\\n\\nRedirecting to document...`);
                        window.location.href = `/documents/${{data.document_id}}`;
                    }} else {{
                        alert('Failed to create document: ' + (data.detail || 'Unknown error'));
                    }}
                }})
                .catch(error => {{
                    alert('Error: ' + error.message);
                }});
            }}
        </script>
    </body>
    </html>
    '''
    
    return html

async def generate_template_prompt(request: AITemplateRequest) -> str:
    """Generate prompt for AI template generation with proper placeholder format"""
    
    template_type_descriptions = {
        "contract": "legally binding agreement between parties",
        "agreement": "formal arrangement between parties",
        "nda": "non-disclosure agreement to protect confidential information",
        "proposal": "formal offer or suggestion",
        "invoice": "bill for products or services",
        "form": "structured document with fields to fill",
        "letter": "formal written communication",
        "resume": "summary of work experience and qualifications",
        "report": "detailed account or statement",
        "memo": "internal business communication",
        "application": "request for consideration",
        "receipt": "proof of payment",
        "quotation": "price estimate",
        "order_form": "document to request products/services",
        "checklist": "list of items to verify",
        "questionnaire": "set of questions for information gathering",
        "feedback_form": "document to collect opinions",
        "survey": "research instrument",
        "certificate": "official document of achievement"
    }
    
    description = request.description
    template_type = request.template_type.value
    style = request.document_style.value
    language = request.language
    tone = request.tone
    
    prompt = f"""Generate a comprehensive {template_type} document template.

DESCRIPTION: {description}
DOCUMENT TYPE: {template_type_descriptions.get(template_type, template_type)}
STYLE: {style}
LANGUAGE: {language}
TONE: {tone}

CRITICAL INSTRUCTIONS FOR FIELD PLACEMENT:
1. Use EXCLUSIVELY DOUBLE CURLY BRACES format for ALL fields: {{field_name}}
2. Field names MUST be descriptive snake_case (e.g., client_name, not client_)
3. NEVER use parentheses (), brackets [], or angle brackets <> for fields
4. Place fields INLINE where data should be entered
5. Include appropriate spacing around fields for readability
6. Make field names clear and self-explanatory

EXAMPLE FORMATS:
CORRECT: "This agreement is between {{client_name}} and {{service_provider_name}}."
INCORRECT: "This agreement is between (client_) and [provider_name]."

STANDARD FIELD CATEGORIES:
- Names: {{first_name}}, {{last_name}}, {{full_name}}, {{company_name}}
- Dates: {{effective_date}}, {{signature_date}}, {{completion_date}}, {{expiration_date}}
- Contact: {{email_address}}, {{phone_number}}, {{mailing_address}}, {{website_url}}
- Financial: {{total_amount}}, {{deposit_amount}}, {{balance_due}}, {{payment_terms}}
- Signatures: {{client_signature}}, {{witness_signature}}, {{notary_signature}}
- Legal: {{jurisdiction}}, {{governing_law}}, {{arbitration_clause}}

DOCUMENT STRUCTURE REQUIREMENTS:
- Professional title and header
- Clear introduction/purpose section
- Well-organized body content with sections
- Terms and conditions where applicable
- Signature blocks with date fields
- Proper footer with legal information

STYLE-SPECIFIC GUIDELINES:
- {style.upper()} STYLE: Use appropriate formatting for this style
- Maintain consistent tone: {tone}
- Ensure proper spacing and alignment

SPECIAL INSTRUCTIONS:"""
    
    if request.fields_to_extract:
        prompt += f"\n\nMUST INCLUDE THESE SPECIFIC FIELDS: {', '.join(request.fields_to_extract)}"
    
    if request.include_clauses:
        prompt += f"\n\nINCLUDE THESE SPECIFIC CLAUSES: {', '.join(request.include_clauses)}"
    
    if request.custom_instructions:
        prompt += f"\n\nADDITIONAL REQUIREMENTS: {request.custom_instructions}"
    
    prompt += f"""

RESPONSE FORMAT (MUST BE VALID JSON):
{{
    "title": "Professional, descriptive title for the template",
    "content": "Full template text with {{field_name}} placeholders. Use ACTUAL standard HTML tags for structure: <h1>, <h2>, <h3>, <p>, <ul>, <li>, <strong>. DO NOT wrap HTML tags in braces - use standard < > brackets.",
    "fields": [
        {{
            "name": "field_name_in_snake_case",
            "label": "Human readable label with proper capitalization",
            "type": "text|date|signature|initial|checkbox|radio|dropdown|email|phone|number|textarea",
            "required": true/false,
            "placeholder": "Helpful example or instruction",
            "description": "Brief description of what this field is for",
            "validation_rules": {{}}
        }}
    ],
    "summary": "Brief 2-3 sentence summary of what this template is for",
    "suggested_tags": ["relevant_tag1", "relevant_tag2", "relevant_tag3"],
    "metadata": {{
        "estimated_completion_time": "10-15 minutes",
        "complexity": "low|medium|high",
        "legal_considerations": ["consideration1", "consideration2"],
        "recommended_use_cases": ["use_case1", "use_case2"]
    }}
}}

FINAL REMINDERS:
1. Use ONLY {{field_name}} format - no other formats allowed
2. Make field names complete and descriptive
3. Ensure proper document structure and flow
4. Include all necessary sections for a {template_type}
5. Add clear instructions where helpful
6. Make it professional and ready for use"""
    
    return prompt

# ========== DATABASE HELPERS ==========

def create_template_document(user_id: str, template_data: Dict, request: AITemplateRequest) -> Dict:
    """Create a template document for MongoDB"""
    template_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    
    content = template_data.get("content", "")
    fields = template_data.get("fields", [])
    
    # Normalize content before processing
    normalized_content, _ = normalize_field_names(content)
    
    # Generate HTML content for both modes
    edit_html = create_edit_mode_html(normalized_content, fields)
    preview_html = create_preview_mode_html(normalized_content, fields, style=request.document_style.value)
    
    template_doc = {
        "_id": template_id,
        "id": template_id,
        "user_id": user_id,
        "title": template_data.get("title", f"Generated {request.template_type.value.title()}"),
        "content": normalized_content,  # Store normalized content
        "original_content": content,     # Keep original for reference
        "edit_html": edit_html,
        "preview_html": preview_html,
        "template_type": request.template_type.value,
        "document_style": request.document_style.value,
        "description": request.description,
        "ai_model_used": "openai",
        "metadata": template_data.get("metadata", {}),
        "summary": template_data.get("summary", ""),
        "suggested_tags": template_data.get("suggested_tags", []),
        "tags": request.tags or [],
        "language": request.language,
        "tone": request.tone,
        "is_active": True,
        "is_public": False,
        "version": 1,
        "created_at": now,
        "updated_at": now,
        "word_count": len(normalized_content.split()),
        "field_count": len(fields),
        "custom_instructions": request.custom_instructions,
        "clauses_included": request.include_clauses or [],
        "normalized": True  # Flag to indicate content has been normalized
    }
    
    return template_doc

def create_field_documents(template_id: str, fields: List[Dict]) -> List[Dict]:
    """Create field documents for MongoDB"""
    field_docs = []
    for field_data in fields:
        field_id = field_data.get("id") or str(uuid.uuid4())
        
        # Calculate default dimensions based on field type
        field_type = field_data.get("field_type") or field_data.get("type", "text")
        width, height = get_default_field_dimensions(field_type)
        
        field_doc = {
            "_id": field_id,
            "id": field_id,
            "template_id": template_id,
            "name": field_data.get("name", ""),
            "label": field_data.get("label", field_data.get("name", "").replace('_', ' ').title()),
            "field_type": field_type,
            "required": field_data.get("required", True),
            "placeholder": field_data.get("placeholder", ""),
            "default_value": field_data.get("default_value"),
            "description": field_data.get("description", ""),
            "options": field_data.get("options", []),
            "x_position": field_data.get("x_position", 10.0),
            "y_position": field_data.get("y_position", 10.0),
            "width": width,
            "height": height,
            "style": field_data.get("style", {}),
            "validation_rules": field_data.get("validation_rules", {}),
            "metadata": field_data.get("metadata", {}),
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }
        field_docs.append(field_doc)
    
    return field_docs

# ========== CORE INTEGRATION FUNCTIONS ==========

def convert_percentage_to_pdf_points(percent_x: float, percent_y: float, 
                                     page_width: float = 612.0, page_height: float = 792.0) -> Dict:
    """Convert percentage-based positions to PDF points"""
    return {
        "x": (percent_x / 100) * page_width,
        "y": (percent_y / 100) * page_height,
        "width": (percent_x / 100) * page_width,
        "height": (percent_y / 100) * page_height
    }

def map_field_type(ai_field_type: str) -> str:
    """Map AI field type to your existing FieldType"""
    field_type_mapping = {
        "text": "textbox",
        "textarea": "textbox",
        "date": "date",
        "signature": "signature",
        "initial": "initials",
        "checkbox": "checkbox",
        "radio": "radio",
        "dropdown": "dropdown",
        "email": "mail",
        "phone": "textbox",
        "number": "textbox",
        "file": "attachment",
        "image": "attachment",
        "qr_code": "textbox"
    }
    return field_type_mapping.get(ai_field_type.lower(), "textbox")

async def convert_template_to_document_internal(template_id: str, user_id: str, 
                                              filename: str = None, envelope_id: str = None) -> Dict:
    """
    Internal function to convert AI template to your core document
    """
    try:
        # Get template
        template = templates_collection.find_one({
            "_id": template_id,
            "user_id": user_id
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get template fields
        template_fields = list(template_fields_collection.find({
            "template_id": template_id
        }))
        
        # Prepare document content (convert to PDF bytes)
        content = template.get("content", "").encode('utf-8')
        
        # Generate a filename if not provided
        if not filename:
            filename = f"{template['title']}.txt"
        
        # Create a mock UploadFile object for your existing upload logic
        from io import BytesIO
        from fastapi import UploadFile
        
        # Create a file-like object
        file_content = BytesIO(content)
        upload_file = UploadFile(
            filename=filename,
            file=file_content,
            content_type="text/plain"
        )
        
        # Prepare envelope_id data
        envelope_id_value = envelope_id
        
        # Import your existing upload function
        from .documents import upload_document
        
        # Call your existing upload endpoint logic
        # Note: This is a simplified version - you'll need to adapt based on your actual implementation
        try:
            # Create document using your existing database logic
            from .documents import fs
            import gridfs
            
            # Convert text to PDF (simplified - use your actual PDF conversion)
            pdf_bytes = content  # In reality, convert to PDF
            
            original_file_id = fs.put(
                content,
                filename=filename,
                content_type="text/plain"
            )
            
            pdf_file_id = fs.put(
                pdf_bytes,
                filename=filename.replace('.txt', '.pdf'),
                content_type="application/pdf"
            )
            
            # Create document record
            doc_data = {
                "filename": filename,
                "uploaded_at": datetime.utcnow(),
                "owner_id": ObjectId(user_id),
                "owner_email": "",  # Will be filled from user lookup
                "mime_type": "text/plain",
                "size": len(content),
                "page_count": 1,
                "status": "draft",
                "common_message": "Please review and sign this document.",
                "original_file_id": original_file_id,
                "pdf_file_id": pdf_file_id,
                "signed_pdf_id": None,
                "recipient_count": 0,
                "signed_count": 0,
                "source": "ai_template",
                "is_converted": True,
                "envelope_id": envelope_id_value,
                "envelope_auto_generated": not bool(envelope_id),
                "template_source_id": template_id  # Link back to template
            }
            
            # Get user email
            user = db.users.find_one({"_id": ObjectId(user_id)})
            if user:
                doc_data["owner_email"] = user.get("email", "")
            
            # Insert document
            result = db.documents.insert_one(doc_data)
            doc_id = result.inserted_id
            
            # Convert template fields to signature fields
            for field in template_fields:
                # Calculate PDF coordinates from percentages
                pdf_coords = convert_percentage_to_pdf_points(
                    field.get("x_position", 10),
                    field.get("y_position", 10)
                )
                
                # Create recipient for this field
                recipient_data = {
                    "document_id": doc_id,
                    "name": "To be assigned",
                    "email": "",
                    "role": "signer",  # Default role
                    "signing_order": 1,
                    "status": "created",
                    "added_at": datetime.utcnow(),
                    "added_by": ObjectId(user_id)
                }
                
                recipient_result = db.recipients.insert_one(recipient_data)
                
                # Create signature field
                signature_field = {
                    "document_id": doc_id,
                    "recipient_id": recipient_result.inserted_id,
                    "type": map_field_type(field.get("field_type", "text")),
                    "page": 0,  # Single page for now
                    "x": pdf_coords["x"],
                    "y": pdf_coords["y"],
                    "width": field.get("width", 30),
                    "height": field.get("height", 5),
                    "required": field.get("required", True),
                    "label": field.get("label", field.get("name", "")),
                    "placeholder": field.get("placeholder", ""),
                    "added_at": datetime.utcnow()
                }
                
                db.signature_fields.insert_one(signature_field)
            
            # Update document with field count
            field_count = len(template_fields)
            db.documents.update_one(
                {"_id": doc_id},
                {"$set": {
                    "field_count": field_count,
                    "recipient_count": field_count  # One recipient per field for now
                }}
            )
            
            # Log the conversion
            db.ai_template_conversions.insert_one({
                "template_id": template_id,
                "document_id": str(doc_id),
                "user_id": user_id,
                "converted_at": datetime.utcnow(),
                "field_count": field_count
            })
            
            return {
                "document_id": str(doc_id),
                "filename": filename,
                "field_count": field_count,
                "template_name": template.get("title", "")
            }
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to create document: {str(e)}"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Template conversion failed: {str(e)}"
        )

# ========== API ENDPOINTS ==========

@router.post("/generate", response_model=AITemplateResponse)
async def generate_ai_template(
    request: AITemplateRequest,
    current_user: dict = Depends(get_current_user)
):
    """Generate a new template using AI"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="User not authenticated")
        
        if not cohere_client:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # Generate prompt
        prompt = await generate_template_prompt(request)
        
        # Call AI
        ai_response = await analyze_with_ai("", prompt, response_format="json_object")
        
        # Validate AI response
        if not isinstance(ai_response, dict):
            raise HTTPException(status_code=500, detail="Invalid AI response format")
        
        # Ensure required fields exist
        ai_response.setdefault("title", f"Generated {request.template_type.value.title()} Template")
        ai_response.setdefault("content", f"# {ai_response['title']}\n\nNo content generated.")
        ai_response.setdefault("fields", [])
        ai_response.setdefault("summary", f"A {request.template_type.value} template")
        ai_response.setdefault("suggested_tags", [])
        ai_response.setdefault("metadata", {})
        
        # Normalize field names in content
        normalized_content, detected_fields = normalize_field_names(ai_response["content"])
        ai_response["content"] = normalized_content
        
        # Create template document
        template_doc = create_template_document(user_id, ai_response, request)
        
        # Create field documents
        fields_data = ai_response.get("fields", [])
        field_docs = create_field_documents(template_doc["_id"], fields_data)
        
        # Save to MongoDB
        try:
            # Insert template
            templates_collection.insert_one(template_doc)
            
            # Insert fields if any
            if field_docs:
                template_fields_collection.insert_many(field_docs)
                
            # Create AI log
            log_doc = {
                "_id": str(uuid.uuid4()),
                "template_id": template_doc["_id"],
                "user_id": user_id,
                "prompt": prompt[:5000],
                "response_summary": {
                    "title": ai_response.get("title", ""),
                    "field_count": len(fields_data),
                    "word_count": len(ai_response.get("content", "").split()),
                    "normalized": True
                },
                "model_used": "openai",
                "created_at": datetime.now(timezone.utc)
            }
            ai_generation_logs_collection.insert_one(log_doc)
            
            # Create initial version
            version_doc = {
                "_id": str(uuid.uuid4()),
                "template_id": template_doc["_id"],
                "version_number": 1,
                "content": template_doc["content"],
                "edit_html": template_doc["edit_html"],
                "preview_html": template_doc["preview_html"],
                "fields": field_docs,
                "created_by": user_id,
                "created_at": datetime.now(timezone.utc),
                "changes": ["Initial version created with normalized field names"]
            }
            template_versions_collection.insert_one(version_doc)
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
        
        # Prepare response
        return AITemplateResponse(
            template_id=template_doc["_id"],
            title=template_doc["title"],
            content=template_doc["content"],
            html_content=template_doc["preview_html"],
            fields=[
                TemplateFieldSchema(
                    id=field.get("id"),
                    name=field.get("name", ""),
                    label=field.get("label", field.get("name", "")),
                    field_type=field.get("field_type") or field.get("type", "text"),
                    required=field.get("required", True),
                    placeholder=field.get("placeholder"),
                    default_value=field.get("default_value"),
                    options=field.get("options", []),
                    x_position=field.get("x_position", 10.0),
                    y_position=field.get("y_position", 10.0),
                    width=field.get("width", 30.0),
                    height=field.get("height", 5.0),
                    style=field.get("style", {}),
                    validation_rules=field.get("validation_rules", {}),
                    metadata=field.get("metadata", {})
                ) for field in field_docs
            ],
            summary=ai_response.get("summary", ""),
            suggested_tags=ai_response.get("suggested_tags", []),
            metadata=ai_response.get("metadata", {}),
            created_at=template_doc["created_at"],
            version=1
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template generation failed: {str(e)}")

@router.get("/templates/{template_id}/edit-mode")
async def get_template_edit_mode(
    template_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get template in edit mode (with placeholders and grid)"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Get template
        template = templates_collection.find_one({
            "_id": template_id,
            "$or": [
                {"user_id": user_id},
                {"is_public": True}
            ]
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get fields
        fields = list(template_fields_collection.find({"template_id": template_id}))
        
        # Check if content needs normalization
        content = template.get("content", "")
        if not template.get("normalized", False):
            # Normalize content if not already done
            normalized_content, _ = normalize_field_names(content)
            
            # Update template with normalized content
            templates_collection.update_one(
                {"_id": template_id},
                {"$set": {
                    "content": normalized_content,
                    "normalized": True,
                    "updated_at": datetime.now(timezone.utc)
                }}
            )
            content = normalized_content
        
        # Generate edit mode HTML if not exists or needs update
        if "edit_html" in template and template["edit_html"] and template.get("normalized", False):
            html_content = template["edit_html"]
        else:
            html_content = create_edit_mode_html(content, fields)
            
            # Update template with generated HTML
            templates_collection.update_one(
                {"_id": template_id},
                {"$set": {
                    "edit_html": html_content,
                    "updated_at": datetime.now(timezone.utc),
                    "normalized": True
                }}
            )
        
        return HTMLResponse(content=html_content)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get edit mode: {str(e)}")

@router.get("/templates/{template_id}/preview-mode")
async def get_template_preview_mode(
    template_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get template in preview mode (with form fields)"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Get template
        template = templates_collection.find_one({
            "_id": template_id,
            "$or": [
                {"user_id": user_id},
                {"is_public": True}
            ]
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get fields
        fields = list(template_fields_collection.find({"template_id": template_id}))
        
        # Check if content needs normalization
        content = template.get("content", "")
        if not template.get("normalized", False):
            # Normalize content if not already done
            normalized_content, _ = normalize_field_names(content)
            content = normalized_content
        
        # Generate preview mode HTML if not exists or needs update
        if "preview_html" in template and template["preview_html"] and template.get("normalized", False):
            html_content = template["preview_html"]
        else:
            html_content = create_preview_mode_html(
                content, 
                fields, 
                style=template.get("document_style", "modern")
            )
            
            # Update template with generated HTML
            templates_collection.update_one(
                {"_id": template_id},
                {"$set": {
                    "preview_html": html_content,
                    "updated_at": datetime.now(timezone.utc),
                    "normalized": True
                }}
            )
        
        return HTMLResponse(content=html_content)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview generation failed: {str(e)}")

@router.post("/analyze-document", response_model=DocumentAnalysisResponse)
async def analyze_document_for_template(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Analyze uploaded document to suggest template structure"""
    try:
        if not cohere_client:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # Extract text from document
        document_text, page_count = await extract_document_text(file)
        
        # Limit text length for analysis
        if len(document_text) > 8000:
            document_text = document_text[:8000] + "... [document truncated]"
        
        # Normalize field names in the document text first
        normalized_text, detected_fields = normalize_field_names(document_text)
        
        # Analyze with OpenAI
        analysis_prompt = f"""Analyze this document and provide a comprehensive analysis for template creation.

DOCUMENT ANALYSIS TASK:
1. Identify the document type (contract, agreement, form, letter, proposal, invoice, etc.)
2. Provide a concise summary
3. Identify all potential form fields users would need to fill
4. Extract key clauses or sections
5. Assess risk level (low, medium, high)
6. Provide recommendations for template creation
7. Suggest the most appropriate template type

DETECTED FIELD PLACEHOLDERS (already normalized):
{", ".join(detected_fields) if detected_fields else "None detected"}

RESPONSE FORMAT (MUST BE VALID JSON):
{{
    "document_type": "contract|agreement|form|letter|proposal|invoice|report|memo|application|other",
    "summary": "Brief 2-3 sentence summary",
    "potential_fields": [
        {{
            "name": "field_name_in_snake_case",
            "label": "Human readable label",
            "type": "text|date|signature|initial|checkbox|radio|dropdown|email|phone|number|textarea",
            "required": true/false,
            "placeholder": "Example or instruction",
            "description": "What this field is for"
        }}
    ],
    "key_clauses": ["clause1", "clause2", "clause3"],
    "risk_level": "low|medium|high",
    "recommendations": ["recommendation1", "recommendation2", "recommendation3"],
    "suggested_template_type": "contract|agreement|form|letter|proposal|invoice|etc"
}}

IMPORTANT: Use standardized field names in snake_case format."""
        
        analysis = await analyze_with_ai(normalized_text, analysis_prompt, response_format="json_object")
        
        # Ensure response has expected structure
        if not isinstance(analysis, dict):
            analysis = {
                "document_type": "other",
                "summary": "Document analysis completed.",
                "potential_fields": [],
                "key_clauses": [],
                "risk_level": "medium",
                "recommendations": ["Review the document manually for specific requirements."],
                "suggested_template_type": "other"
            }
        
        return DocumentAnalysisResponse(
            document_type=analysis.get("document_type", "other"),
            summary=analysis.get("summary", ""),
            potential_fields=[
                TemplateFieldSchema(
                    name=field.get("name", f"field_{i}"),
                    label=field.get("label", field.get("name", f"Field {i}")),
                    field_type=field.get("type", "text"),
                    required=field.get("required", True),
                    placeholder=field.get("placeholder", ""),
                    description=field.get("description", "")
                ) for i, field in enumerate(analysis.get("potential_fields", []))
            ],
            key_clauses=analysis.get("key_clauses", []),
            risk_level=analysis.get("risk_level", "medium"),
            recommendations=analysis.get("recommendations", []),
            word_count=len(document_text.split()),
            page_count=page_count,
            suggested_template_type=TemplateType(analysis.get("suggested_template_type", "other"))
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document analysis failed: {str(e)}")

@router.post("/templates/{template_id}/convert-to-document")
async def convert_template_to_document_endpoint(
    request: TemplateToDocumentRequest,
    current_user: dict = Depends(get_current_user)
):
    """Convert AI template to a core document ready for signing"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="User not authenticated")
        
        # Convert template to document
        result = await convert_template_to_document_internal(
            request.template_id,
            user_id,
            request.filename,
            request.envelope_id
        )
        
        return JSONResponse(content={
            "success": True,
            "message": "Template converted to document successfully",
            "document_id": result["document_id"],
            "filename": result["filename"],
            "field_count": result["field_count"],
            "template_name": result["template_name"]
        })
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

@router.post("/templates/{template_id}/normalize")
async def normalize_template(
    template_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Normalize field names in an existing template"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Get template
        template = templates_collection.find_one({
            "_id": template_id,
            "user_id": user_id
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get fields
        fields = list(template_fields_collection.find({"template_id": template_id}))
        
        # Normalize content
        content = template.get("content", "")
        normalized_content, detected_fields = normalize_field_names(content)
        
        # Update fields if needed
        updated_field_names = set(detected_fields)
        existing_field_names = {field.get("name", "") for field in fields}
        
        # Add missing fields
        missing_fields = updated_field_names - existing_field_names
        for field_name in missing_fields:
            # Create new field
            field_doc = {
                "_id": str(uuid.uuid4()),
                "id": str(uuid.uuid4()),
                "template_id": template_id,
                "name": field_name,
                "label": field_name.replace('_', ' ').title(),
                "field_type": "text",
                "required": True,
                "placeholder": f"Enter {field_name.replace('_', ' ')}",
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }
            template_fields_collection.insert_one(field_doc)
            fields.append(field_doc)
        
        # Regenerate HTML
        edit_html = create_edit_mode_html(normalized_content, fields)
        preview_html = create_preview_mode_html(
            normalized_content, 
            fields, 
            style=template.get("document_style", "modern")
        )
        
        # Update template
        update_data = {
            "content": normalized_content,
            "original_content": content,
            "edit_html": edit_html,
            "preview_html": preview_html,
            "normalized": True,
            "updated_at": datetime.now(timezone.utc),
            "word_count": len(normalized_content.split()),
            "field_count": len(fields)
        }
        
        templates_collection.update_one(
            {"_id": template_id},
            {"$set": update_data}
        )
        
        # Create version record
        version_doc = {
            "_id": str(uuid.uuid4()),
            "template_id": template_id,
            "version_number": template.get("version", 1) + 1,
            "content": normalized_content,
            "edit_html": edit_html,
            "preview_html": preview_html,
            "fields": fields,
            "created_by": user_id,
            "created_at": datetime.now(timezone.utc),
            "changes": ["Normalized field names and regenerated HTML"]
        }
        template_versions_collection.insert_one(version_doc)
        
        # Update template version
        templates_collection.update_one(
            {"_id": template_id},
            {"$inc": {"version": 1}}
        )
        
        return {
            "success": True,
            "message": "Template normalized successfully",
            "detected_fields": list(detected_fields),
            "missing_fields_added": len(missing_fields),
            "new_version": template.get("version", 1) + 1
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Normalization failed: {str(e)}")

@router.get("/templates/{template_id}")
async def get_template(
    template_id: str,
    mode: str = Query("preview", pattern="^(edit|preview)$"),
    current_user: dict = Depends(get_current_user)
):
    """Get template in specified mode (edit or preview)"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Get template
        template = templates_collection.find_one({
            "_id": template_id,
            "$or": [
                {"user_id": user_id},
                {"is_public": True}
            ]
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get fields
        fields = list(template_fields_collection.find({"template_id": template_id}))
        
        # Ensure content is normalized
        content = template.get("content", "")
        if not template.get("normalized", False):
            normalized_content, _ = normalize_field_names(content)
            content = normalized_content
        
        if mode == "edit":
            # Generate/edit mode HTML
            if "edit_html" in template and template["edit_html"] and template.get("normalized", False):
                html_content = template["edit_html"]
            else:
                html_content = create_edit_mode_html(content, fields)
        else:
            # Preview mode HTML
            if "preview_html" in template and template["preview_html"] and template.get("normalized", False):
                html_content = template["preview_html"]
            else:
                html_content = create_preview_mode_html(
                    content, 
                    fields, 
                    style=template.get("document_style", "modern")
                )
        
        return HTMLResponse(content=html_content)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get template: {str(e)}")

@router.put("/templates/{template_id}")
async def update_template(
    template_id: str,
    content: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """Update template content"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Verify template ownership
        template = templates_collection.find_one({
            "_id": template_id,
            "user_id": user_id
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Get fields
        fields = list(template_fields_collection.find({"template_id": template_id}))
        
        # Normalize the new content
        normalized_content, detected_fields = normalize_field_names(content)
        
        # Regenerate HTML for both modes
        edit_html = create_edit_mode_html(normalized_content, fields)
        preview_html = create_preview_mode_html(normalized_content, fields, style=template.get("document_style", "modern"))
        
        # Update template
        update_data = {
            "content": normalized_content,
            "original_content": content,
            "edit_html": edit_html,
            "preview_html": preview_html,
            "updated_at": datetime.now(timezone.utc),
            "word_count": len(normalized_content.split()),
            "version": template.get("version", 1) + 1,
            "normalized": True
        }
        
        templates_collection.update_one(
            {"_id": template_id},
            {"$set": update_data}
        )
        
        # Create version record
        version_doc = {
            "_id": str(uuid.uuid4()),
            "template_id": template_id,
            "version_number": update_data["version"],
            "content": normalized_content,
            "edit_html": edit_html,
            "preview_html": preview_html,
            "fields": fields,
            "created_by": user_id,
            "created_at": datetime.now(timezone.utc),
            "changes": ["Content updated with normalized field names"]
        }
        template_versions_collection.insert_one(version_doc)
        
        return {
            "success": True,
            "message": "Template updated successfully",
            "version": update_data["version"],
            "detected_fields": list(detected_fields)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update template: {str(e)}")

@router.post("/templates/{template_id}/fields")
async def add_template_field(
    template_id: str,
    field: TemplateFieldSchema,
    current_user: dict = Depends(get_current_user)
):
    """Add a field to a template"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Verify template ownership
        template = templates_collection.find_one({
            "_id": template_id,
            "user_id": user_id
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Create field document
        field_id = str(uuid.uuid4())
        field_doc = {
            "_id": field_id,
            "id": field_id,
            "template_id": template_id,
            "name": field.name,
            "label": field.label,
            "field_type": field.field_type,
            "required": field.required,
            "placeholder": field.placeholder,
            "default_value": field.default_value,
            "options": field.options or [],
            "x_position": field.x_position,
            "y_position": field.y_position,
            "width": field.width,
            "height": field.height,
            "style": field.style or {},
            "validation_rules": field.validation_rules or {},
            "metadata": field.metadata or {},
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }
        
        # Insert field
        template_fields_collection.insert_one(field_doc)
        
        # Update template to refresh HTML
        fields = list(template_fields_collection.find({"template_id": template_id}))
        content = template.get("content", "")
        
        # Add placeholder to content if not already present
        placeholder = f"{{{{{field.name}}}}}"
        if placeholder not in content:
            # Add placeholder at the end of the content
            content += f"\n\n{placeholder}"
            
            # Normalize the updated content
            normalized_content, _ = normalize_field_names(content)
            
            # Update template content
            templates_collection.update_one(
                {"_id": template_id},
                {"$set": {"content": normalized_content, "original_content": content}}
            )
            content = normalized_content
        
        # Regenerate HTML
        edit_html = create_edit_mode_html(content, fields)
        preview_html = create_preview_mode_html(content, fields, style=template.get("document_style", "modern"))
        
        templates_collection.update_one(
            {"_id": template_id},
            {"$set": {
                "edit_html": edit_html,
                "preview_html": preview_html,
                "updated_at": datetime.now(timezone.utc),
                "field_count": len(fields),
                "normalized": True
            }}
        )
        
        return {
            "success": True,
            "field_id": field_id,
            "message": f"Field '{field.name}' added successfully",
            "placeholder_format": "{{field_name}}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to add field: {str(e)}")

@router.get("/template-types")
async def get_template_types():
    """Get all available template types with descriptions"""
    template_types = [
        {
            "value": TemplateType.CONTRACT.value,
            "label": "Contract",
            "description": "Legally binding agreement between parties",
            "icon": "📝",
            "common_uses": ["Service agreements", "Employment contracts", "Sales contracts"]
        },
        {
            "value": TemplateType.AGREEMENT.value,
            "label": "Agreement",
            "description": "Formal arrangement between parties",
            "icon": "🤝",
            "common_uses": ["Partnership agreements", "Rental agreements", "License agreements"]
        },
        {
            "value": TemplateType.NDA.value,
            "label": "Non-Disclosure Agreement",
            "description": "Protects confidential information",
            "icon": "🔒",
            "common_uses": ["Confidentiality agreements", "Trade secret protection"]
        },
        {
            "value": TemplateType.PROPOSAL.value,
            "label": "Proposal",
            "description": "Formal offer or suggestion",
            "icon": "📄",
            "common_uses": ["Business proposals", "Project proposals", "Grant proposals"]
        },
        {
            "value": TemplateType.INVOICE.value,
            "label": "Invoice",
            "description": "Bill for products or services",
            "icon": "🧾",
            "common_uses": ["Service billing", "Product sales", "Freelance work"]
        },
        {
            "value": TemplateType.FORM.value,
            "label": "Form",
            "description": "Structured document with fields to fill",
            "icon": "📋",
            "common_uses": ["Application forms", "Registration forms", "Survey forms"]
        },
        {
            "value": TemplateType.LETTER.value,
            "label": "Letter",
            "description": "Formal written communication",
            "icon": "✉️",
            "common_uses": ["Business letters", "Cover letters", "Official correspondence"]
        },
        {
            "value": TemplateType.RESUME.value,
            "label": "Resume/CV",
            "description": "Summary of work experience and qualifications",
            "icon": "📄",
            "common_uses": ["Job applications", "Professional profiles"]
        },
        {
            "value": TemplateType.REPORT.value,
            "label": "Report",
            "description": "Detailed account or statement",
            "icon": "📊",
            "common_uses": ["Business reports", "Research reports", "Progress reports"]
        },
        {
            "value": TemplateType.MEMO.value,
            "label": "Memo",
            "description": "Internal business communication",
            "icon": "📝",
            "common_uses": ["Internal announcements", "Policy updates", "Meeting minutes"]
        }
    ]
    
    return template_types

@router.get("/user-templates")
async def get_user_templates(
    current_user: dict = Depends(get_current_user),
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=100),
    template_type: Optional[str] = Query(None)
):
    """Get all templates for the current user"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        query = {"user_id": user_id, "is_active": True}
        if template_type:
            query["template_type"] = template_type
        
        templates = list(templates_collection.find(query)
                         .sort("created_at", -1)
                         .skip(skip)
                         .limit(limit))
        
        # Format response
        template_list = []
        for template in templates:
            template_list.append({
                "id": template.get("_id"),
                "title": template.get("title"),
                "template_type": template.get("template_type"),
                "description": template.get("description"),
                "created_at": template.get("created_at"),
                "updated_at": template.get("updated_at"),
                "field_count": template.get("field_count", 0),
                "word_count": template.get("word_count", 0),
                "version": template.get("version", 1),
                "is_public": template.get("is_public", False)
            })
        
        total_count = templates_collection.count_documents(query)
        
        return {
            "templates": template_list,
            "total": total_count,
            "page": (skip // limit) + 1,
            "page_size": limit,
            "total_pages": (total_count + limit - 1) // limit
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get templates: {str(e)}")

@router.delete("/templates/{template_id}")
async def delete_template(
    template_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Soft delete a template"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Verify template ownership
        template = templates_collection.find_one({
            "_id": template_id,
            "user_id": user_id
        })
        
        if not template:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Soft delete
        templates_collection.update_one(
            {"_id": template_id},
            {"$set": {
                "is_active": False,
                "deleted_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }}
        )
        
        return {
            "success": True,
            "message": "Template deleted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete template: {str(e)}")

@router.post("/templates/{template_id}/clone")
async def clone_template(
    template_id: str,
    new_name: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """Clone an existing template"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Get original template
        original = templates_collection.find_one({
            "_id": template_id,
            "$or": [
                {"user_id": user_id},
                {"is_public": True}
            ]
        })
        
        if not original:
            raise HTTPException(status_code=404, detail="Template not found")
        
        # Create new template ID
        new_template_id = str(uuid.uuid4())
        
        # Clone template document
        new_template = original.copy()
        new_template["_id"] = new_template_id
        new_template["id"] = new_template_id
        new_template["user_id"] = user_id
        new_template["title"] = new_name
        new_template["is_public"] = False
        new_template["version"] = 1
        new_template["created_at"] = datetime.now(timezone.utc)
        new_template["updated_at"] = datetime.now(timezone.utc)
        new_template.pop("_id", None)  # Remove original ID
        
        # Insert new template
        templates_collection.insert_one(new_template)
        
        # Clone fields
        original_fields = list(template_fields_collection.find({"template_id": template_id}))
        if original_fields:
            new_fields = []
            for field in original_fields:
                new_field = field.copy()
                new_field_id = str(uuid.uuid4())
                new_field["_id"] = new_field_id
                new_field["id"] = new_field_id
                new_field["template_id"] = new_template_id
                new_field["created_at"] = datetime.now(timezone.utc)
                new_field["updated_at"] = datetime.now(timezone.utc)
                new_field.pop("_id", None)  # Remove original ID
                new_fields.append(new_field)
            
            template_fields_collection.insert_many(new_fields)
        
        # Create initial version
        version_doc = {
            "_id": str(uuid.uuid4()),
            "template_id": new_template_id,
            "version_number": 1,
            "content": new_template.get("content", ""),
            "edit_html": new_template.get("edit_html", ""),
            "preview_html": new_template.get("preview_html", ""),
            "fields": new_fields,
            "created_by": user_id,
            "created_at": datetime.now(timezone.utc),
            "changes": ["Cloned from template"]
        }
        template_versions_collection.insert_one(version_doc)
        
        return {
            "success": True,
            "new_template_id": new_template_id,
            "message": f"Template cloned as '{new_name}'"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Clone failed: {str(e)}")

@router.get("/statistics")
async def get_template_statistics(
    current_user: dict = Depends(get_current_user)
):
    """Get template usage statistics for the current user"""
    try:
        user_id = current_user.get("id") or current_user.get("_id")
        
        # Count templates by type
        pipeline = [
            {"$match": {"user_id": user_id, "is_active": True}},
            {"$group": {"_id": "$template_type", "count": {"$sum": 1}}}
        ]
        
        type_counts = list(templates_collection.aggregate(pipeline))
        
        # Total templates
        total_templates = templates_collection.count_documents({
            "user_id": user_id,
            "is_active": True
        })
        
        # Recent activity
        recent_templates = list(templates_collection.find(
            {"user_id": user_id, "is_active": True},
            {"title": 1, "template_type": 1, "updated_at": 1, "field_count": 1}
        ).sort("updated_at", -1).limit(5))
        
        # Field statistics
        field_pipeline = [
            {"$match": {"template_id": {"$in": [t["_id"] for t in recent_templates]}}},
            {"$group": {"_id": "$field_type", "count": {"$sum": 1}}}
        ]
        
        field_counts = list(template_fields_collection.aggregate(field_pipeline))
        
        return {
            "total_templates": total_templates,
            "templates_by_type": type_counts,
            "recent_templates": recent_templates,
            "field_statistics": field_counts,
            "ai_generations": ai_generation_logs_collection.count_documents({"user_id": user_id})
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get statistics: {str(e)}")

@router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "openai_configured": openai_client is not None,
        "mongodb_connected": templates_collection is not None,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "features": {
            "field_normalization": "enabled",
            "edit_mode": "enhanced",
            "preview_mode": "professional",
            "responsive_design": "enabled"
        }
    }

# ========== SMART FIELD SUGGESTION ENDPOINTS ==========

@router.post("/smart-fields")
async def suggest_smart_fields(
    template_content: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """
    Analyze template content and suggest smart field placements with positions.
    Uses OpenAI to detect places where user input fields should be added.
    """
    try:
        if not cohere_client:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")

        prompt = f"""
Analyze this document template and identify where form fields should be placed.

DOCUMENT CONTENT:
{template_content[:4000]}

TASKS:
1. Detect every part of the document where a user must enter data.
2. For each detected field, provide:
   - name: snake_case identifier (e.g., "client_name")
   - label: human-friendly label (e.g., "Client Name")
   - type: one of [text, textarea, number, email, phone, date, signature, initial, checkbox, dropdown]
   - required: true/false
   - placeholder: helpful user hint
   - description: what this field is for
3. ALSO estimate position coordinates where the field should appear:
   - x_position: horizontal position (0-100 percentage of document width)
   - y_position: vertical position (0-100 percentage of document height)
   - width: field width as percentage of document width (10-50)
   - height: field height as percentage of document height (2-20)

3. Provide improvement suggestions for the document structure.

POSITIONING GUIDELINES:
- First field (like title) should be near top (y_position: 5-10)
- Inline fields (like names in sentences) should be placed where text would flow
- Signature fields go near bottom (y_position: 80-90)
- Make fields proportionate to expected content length
- Leave space between fields (at least 5% margin)

RESPONSE FORMAT (STRICT JSON):
{{
  "fields": [
    {{
      "name": "client_name",
      "label": "Client Name",
      "type": "text",
      "required": true,
      "placeholder": "Enter full legal name",
      "description": "The legal name of the client",
      "x_position": 10.5,
      "y_position": 15.2,
      "width": 30.0,
      "height": 4.0
    }}
  ],
  "improvements": ["Add a clear title", "Include signature section"],
  "estimated_completion_time": "5 minutes",
  "position_analysis": "Fields placed based on document flow: names at top, signatures at bottom"
}}

IMPORTANT: Fields should be placed INLINE with the text using {{field_name}} format.
Calculate positions based on logical document flow."""

        ai_result = await analyze_with_ai("", prompt, response_format="json_object")

        # Safety fallback
        if not isinstance(ai_result, dict):
            ai_result = {
                "fields": [],
                "improvements": [],
                "estimated_completion_time": "Unknown",
                "position_analysis": "Could not analyze positions"
            }

        return JSONResponse(content=ai_result)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Smart field suggestion failed: {str(e)}")

@router.post("/detect-and-position-placeholders")
async def detect_and_position_placeholders(
    template_content: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """
    Detect placeholders in template content and suggest field positions.
    """
    try:
        if not cohere_client:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # First, detect placeholders in content
        placeholder_patterns = [
            r'\{\{([^{}]+)\}\}',  # {{field_name}}
            r'\{([^{}]+)\}',      # {field_name}
            r'\(([^()]+)\)',      # (field_name)
            r'\[([^\[\]]+)\]',    # [field_name]
            r'<([^<>]+)>',        # <field_name>
        ]
        
        detected_fields = []
        for pattern in placeholder_patterns:
            matches = re.findall(pattern, template_content)
            for match in matches:
                field_name = match.strip()
                if field_name and field_name not in [f['name'] for f in detected_fields]:
                    detected_fields.append({
                        "name": field_name,
                        "label": field_name.replace('_', ' ').title(),
                        "type": "text",  # Default type
                        "required": True,
                        "placeholder": f"Enter {field_name.replace('_', ' ')}",
                        "detected_by": pattern
                    })
        
        if not detected_fields:
            return JSONResponse(content={
                "message": "No placeholders detected in content",
                "detected_fields": [],
                "suggested_fields": []
            })
        
        # Now get AI to suggest types and positions
        field_list_text = "\n".join([f"- {field['name']}" for field in detected_fields])
        
        prompt = f"""
Analyze this document template and suggest field types and positions for detected placeholders.

DOCUMENT CONTENT:
{template_content[:4000]}

DETECTED PLACEHOLDERS:
{field_list_text}

TASK:
1. Suggest appropriate field types for each placeholder
   Common types: text, textarea, number, email, phone, date, signature, initial, checkbox, dropdown
2. Calculate optimal positions (x, y, width, height as percentages 0-100)
3. Suggest if field should be required or optional

ANALYSIS GUIDELINES:
- Names and titles: text (25-40% width)
- Dates: date (20-30% width)
- Email addresses: email (30-45% width)
- Phone numbers: phone (30-40% width)
- Amounts: number (20-35% width)
- Addresses: textarea or text (40-60% width)
- Signatures: signature (50-70% width, 15-25% height)
- Descriptions: textarea (60-80% width, 10-20% height)

POSITION CALCULATION:
- Scan where each placeholder appears in text
- Position fields at logical insertion points
- Maintain document flow
- No overlapping fields

RESPONSE FORMAT (STRICT JSON):
{{
  "suggested_fields": [
    {{
      "name": "field_name",
      "type": "text|date|signature|etc",
      "required": true/false,
      "placeholder": "User hint text",
      "x_position": 15.0,
      "y_position": 25.5,
      "width": 35.0,
      "height": 5.0,
      "reason": "Positioned after 'Client:' label in first paragraph"
    }}
  ],
  "detection_summary": "Found X placeholders, suggested positions based on content flow"
}}"""

        ai_result = await analyze_with_ai("", prompt, response_format="json_object")
        
        if not isinstance(ai_result, dict) or "suggested_fields" not in ai_result:
            # Fallback: assign default positions
            suggested_fields = []
            for i, field in enumerate(detected_fields):
                suggested_fields.append({
                    **field,
                    "x_position": 10.0,
                    "y_position": 10.0 + (i * 8),
                    "width": 30.0,
                    "height": 5.0,
                    "reason": "Default sequential positioning"
                })
            
            ai_result = {
                "suggested_fields": suggested_fields,
                "detection_summary": f"Found {len(detected_fields)} placeholders, used default positioning"
            }
        
        return JSONResponse(content={
            "detected_fields": detected_fields,
            **ai_result
        })
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Placeholder detection failed: {str(e)}")

@router.post("/auto-position-fields")
async def auto_position_fields(
    template_content: str = Form(...),
    fields_data: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """
    Automatically calculate optimal positions for fields based on template content.
    """
    try:
        if not cohere_client:
            raise HTTPException(status_code=500, detail="OpenAI API key not configured")
        
        # Parse fields data
        try:
            fields = json.loads(fields_data)
            if not isinstance(fields, list):
                fields = []
        except:
            fields = []
        
        if not fields:
            return JSONResponse(content={"fields": [], "analysis": "No fields to position"})
        
        # Prepare field information for AI
        field_descriptions = "\n".join([
            f"{i+1}. {field.get('name', f'field_{i}')} - {field.get('label', 'No label')} "
            f"(Type: {field.get('type', 'text')})"
            for i, field in enumerate(fields)
        ])
        
        prompt = f"""
Analyze this document template and calculate optimal positions for form fields.

DOCUMENT CONTENT:
{template_content[:4000]}

FIELDS TO POSITION:
{field_descriptions}

TASK:
Calculate optimal x_position, y_position, width, and height for each field.
All values are percentages (0-100).

POSITIONING RULES:
1. Scan document for natural insertion points for each field
2. Place fields where they logically fit in the document flow
3. First fields (titles, names) go near top (y: 5-20)
4. Middle fields (descriptions, amounts) go in middle (y: 25-70)
5. Final fields (signatures, dates) go near bottom (y: 75-90)
6. Inline fields should align with text (x: based on line position)
7. Make width appropriate for content (names: 25-40%, addresses: 40-60%, signatures: 50-70%)
8. Standard height: 4-8% for text fields, 8-15% for textareas, 15-25% for signatures

FIELD TYPES AND TYPICAL DIMENSIONS:
- text: height 4-6%, width 25-40%
- email/phone: height 4-6%, width 30-45%
- date: height 4-6%, width 20-30%
- textarea: height 10-20%, width 50-80%
- signature: height 15-25%, width 50-70%
- dropdown: height 4-8%, width 25-40%

RESPONSE FORMAT (STRICT JSON):
{{
  "positioned_fields": [
    {{
      "name": "field_name",
      "x_position": 12.5,
      "y_position": 18.3,
      "width": 35.0,
      "height": 5.0,
      "reason": "Placed after 'Client Name:' label at beginning of document"
    }}
  ],
  "layout_summary": "Fields positioned for optimal document flow",
  "grid_suggested": true,
  "recommended_grid_size": 5.0
}}

IMPORTANT: Calculate based on document structure. Fields should not overlap.
Leave at least 2% spacing between fields."""

        ai_result = await analyze_with_ai("", prompt, response_format="json_object")
        
        # Merge AI positions with original field data
        if isinstance(ai_result, dict) and "positioned_fields" in ai_result:
            positioned_fields = ai_result["positioned_fields"]
            
            # Create mapping of field positions
            position_map = {field["name"]: field for field in positioned_fields}
            
            # Update original fields with positions
            updated_fields = []
            for field in fields:
                field_name = field.get("name")
                if field_name in position_map:
                    position_data = position_map[field_name]
                    updated_field = field.copy()
                    updated_field.update({
                        "x_position": position_data.get("x_position", 10.0),
                        "y_position": position_data.get("y_position", 10.0),
                        "width": position_data.get("width", 30.0),
                        "height": position_data.get("height", 5.0),
                        "position_reason": position_data.get("reason", "")
                    })
                    updated_fields.append(updated_field)
                else:
                    # Assign default position
                    updated_field = field.copy()
                    updated_field.update({
                        "x_position": 10.0,
                        "y_position": 10.0 + (len(updated_fields) * 10),
                        "width": 30.0,
                        "height": 5.0,
                        "position_reason": "Default position"
                    })
                    updated_fields.append(updated_field)
            
            ai_result["fields"] = updated_fields
        else:
            # Assign sequential positions as fallback
            updated_fields = []
            for i, field in enumerate(fields):
                updated_field = field.copy()
                updated_field.update({
                    "x_position": 10.0,
                    "y_position": 10.0 + (i * 8),
                    "width": 30.0,
                    "height": 5.0,
                    "position_reason": "Sequential positioning"
                })
                updated_fields.append(field)
            
            ai_result = {
                "fields": updated_fields,
                "layout_summary": "Fields positioned sequentially",
                "grid_suggested": False,
                "recommended_grid_size": 5.0
            }
        
        return JSONResponse(content=ai_result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Auto-positioning failed: {str(e)}")

# ========== MAIN FUNCTION FOR TESTING ==========
if __name__ == "__main__":
    print("AI Template Builder Module Loaded Successfully")
    print(f"Cohere Configured: {cohere_client is not None}")
    print(f"Available Endpoints: {[route.path for route in router.routes if hasattr(route, 'path')]}")